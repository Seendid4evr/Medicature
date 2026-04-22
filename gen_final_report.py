"""
Rebuild Medicure_Final_Report.docx from scratch with the WBS image properly embedded.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WBS_IMG = r'c:\xampp\htdocs\medicure\wbs_diagram.png'
OUT     = r'c:\xampp\htdocs\medicure\Medicure_Final_Report.docx'

def rgb(h):
    h=h.lstrip('#'); return RGBColor(int(h[0:2],16),int(h[2:4],16),int(h[4:6],16))

def cell_bg(cell, h):
    tc=cell._tc; p=tc.get_or_add_tcPr()
    s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),h.lstrip('#'))
    p.append(s)

def cell_border(cell, color='CCCCCC', sz='4'):
    tc=cell._tc; p=tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        e=OxmlElement(f'w:{side}')
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),sz)
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),color)
        b.append(e)
    p.append(b)

def H(doc, text, size=16, color='1A3C5E', sb=10, sa=4):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=rgb(color)

def B(doc, text, size=10.5, italic=False, sa=4):
    p=doc.add_paragraph(text); p.paragraph_format.space_after=Pt(sa)
    for r in p.runs:
        r.font.size=Pt(size)
        if italic: r.italic=True

def T(doc, headers, rows, widths=None, hbg='1A3C5E', alt='EBF5FF'):
    t=doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hr=t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; cell_bg(c,hbg); cell_border(c,'FFFFFF','6')
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.color.rgb=rgb('FFFFFF'); r.font.size=Pt(9.5)
    for ri,row in enumerate(rows):
        tr=t.rows[ri+1]; bg=alt if ri%2 else 'FFFFFF'
        for ci,val in enumerate(row):
            c=tr.cells[ci]; cell_bg(c,bg); cell_border(c,'CCCCCC','4')
            p=c.paragraphs[0]; v=str(val)
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT
            r=p.add_run(v); r.font.size=Pt(9.5)
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph()

# ── Build ──────────────────────────────────────────────────────────────────────
doc=Document()
sec=doc.sections[0]
sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
doc.styles['Normal'].font.name='Calibri'
doc.styles['Normal'].font.size=Pt(10.5)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before=Pt(20)
r=p.add_run('MEDICURE'); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=rgb('1A3C5E')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Digital Health & Pharmacy Management Platform')
r.bold=True; r.font.size=Pt(14); r.font.color.rgb=rgb('2563EB')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Final Project Report'); r.font.size=Pt(12); r.font.color.rgb=rgb('555555')
doc.add_paragraph()

it=doc.add_table(rows=5,cols=2); it.style='Table Grid'; it.alignment=WD_TABLE_ALIGNMENT.CENTER
info=[('Prepared By','Saleh Kabir (Seendid)'),
      ('Platform','Web Application - PHP / MySQL / XAMPP'),
      ('Team','2 Developers (1 Senior, 1 Junior)'),
      ('Duration','3 Months  |  April - June 2026'),
      ('Repository','github.com/Seendid4evr/Medicature')]
for i,(k,v) in enumerate(info):
    row=it.rows[i]
    cell_bg(row.cells[0],'1A3C5E'); cell_bg(row.cells[1],'EBF5FF' if i%2 else 'FFFFFF')
    cell_border(row.cells[0],'FFFFFF','4'); cell_border(row.cells[1],'CCCCCC','4')
    rk=row.cells[0].paragraphs[0].add_run(k)
    rk.bold=True; rk.font.color.rgb=rgb('FFFFFF'); rk.font.size=Pt(10)
    rv=row.cells[1].paragraphs[0].add_run(v); rv.font.size=Pt(10)
for row in it.rows:
    row.cells[0].width=Inches(2.0); row.cells[1].width=Inches(4.5)
doc.add_paragraph()

# 1. Project Description
H(doc,'1. Project Description')
B(doc,'Medicure is a comprehensive, full-stack digital health platform built to address medication management challenges faced by patients and caregivers in Bangladesh. It combines medication scheduling, AI-powered symptom checking, a drug reference database of over 21,000 Bangladeshi medicines, and an integrated e-pharmacy system - all in a single mobile-responsive web application built with PHP, MySQL, HTML/CSS/JavaScript, and XAMPP.')
H(doc,'Core Features', size=12, color='2563EB', sb=6)
features=['Smart medication tracking with browser-based alarms',
          'AI-powered symptom triage with OTC medicine suggestions and mandatory safety disclaimer',
          'Searchable database of 21,000+ Bangladeshi brand medicines',
          'Family member / dependent health management',
          'E-pharmacy integration - Buy prescriptions from Arogga',
          'Printable 7-day medication adherence reports',
          'BMI and Ideal Weight health calculators',
          'Admin dashboard: manage patients, send notifications, view full profiles',
          'Password recovery, notification settings, profile management',
          'Progressive Web App (PWA) support for mobile installation']
for f in features:
    p=doc.add_paragraph(f, style='List Bullet')
    p.runs[0].font.size=Pt(10); p.paragraph_format.space_after=Pt(2)
doc.add_paragraph()

# 2. Stakeholders
H(doc,'2. Stakeholders')
T(doc,['#','Stakeholder','Role','Interest'],
  [('1','Patients (Primary Users)','End-users','Track medications, reminders, drug info, buy medicines'),
   ('2','Family Caregivers','Secondary users','Manage medications for dependents'),
   ('3','System Administrator','Platform manager','Oversee users, send notifications'),
   ('4','Pharmacies / Arogga','External partner','Receive purchase redirects'),
   ('5','Healthcare Professionals','Indirect beneficiaries','Receive patient adherence reports'),
   ('6','Developer / Project Owner','Builder & maintainer','Development, deployment, maintenance'),
   ('7','University / Assessors','Evaluators','Review as academic submission')],
  widths=[0.3,1.6,1.5,2.9])

# 3. Value Proposition
H(doc,'3. Value Proposition')
p=doc.add_paragraph()
r=p.add_run('Medicure gives every Bangladeshi patient a personal, intelligent, and connected health management system - for free.')
r.italic=True; r.bold=True; r.font.size=Pt(11); r.font.color.rgb=rgb('2563EB')
doc.add_paragraph()
T(doc,['Problem','Medicure Solution'],
  [('Patients forget medicines','Browser alarm fires at exact scheduled time'),
   ('Do not know what medicine does','21,000+ drug database with generics, side effects, indications'),
   ('Managing family medicines','Dependent system: one account manages all family members'),
   ('Refilling prescriptions','Buy Now button links directly to Arogga with medicine pre-searched'),
   ('No medication history','7-day adherence report, printable for doctors'),
   ('Uncertain about symptoms','AI Triage suggests OTC medicines with mandatory safety disclaimer'),
   ('Admin has no visibility','Full admin dashboard with patient details and history')],
  widths=[2.5,4.3])

# 4. WBS
H(doc,'4. Work Breakdown Structure (WBS)')
B(doc,'The project is decomposed using a Phase-Based WBS into 5 phases and 18 work packages across 3 months.')

# ── IMAGE (native python-docx - proper embedding) ──────────────────────────────
pic_p = doc.add_paragraph()
pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pic_p.add_run().add_picture(WBS_IMG, width=Inches(6.6))

cap_p = doc.add_paragraph()
cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_r = cap_p.add_run('Figure 1: Medicure WBS - Phase-Based | 5 Phases | 18 Work Packages | 708 Total Hours')
cap_r.italic=True; cap_r.font.size=Pt(9); cap_r.font.color.rgb=rgb('555555')
doc.add_paragraph()
# ───────────────────────────────────────────────────────────────────────────────

T(doc,['WBS ID','Phase / Task','Assigned To','Hours','Month'],
  [('1.1','Requirements & Scope Analysis','Both Devs','20','Month 1'),
   ('1.2','Feasibility Study','Dev 1 (Senior)','12','Month 1'),
   ('1.3','Project Charter & Planning','Both Devs','12','Month 1'),
   ('2.1','UI/UX Wireframes & Mockups','Dev 2 (Junior)','32','Month 1'),
   ('2.2','Database Schema Design','Dev 1 (Senior)','24','Month 1'),
   ('2.3','System Architecture Planning','Dev 1 (Senior)','16','Month 1'),
   ('3.1','Backend - PHP, Auth, DB, APIs','Dev 1 (Senior)','136','Month 2'),
   ('3.2','Frontend - HTML/CSS/JS, Dashboard','Dev 2 (Junior)','112','Month 2'),
   ('3.3','Symptom Checker AI Module','Dev 1 (Senior)','56','Month 2'),
   ('3.4','Pharmacy & Medicines Module','Both Devs','56','Month 2'),
   ('3.5','Family Health Management Module','Both Devs','40','Month 2'),
   ('3.6','Reports & Health Calculators','Both Devs','40','Month 2'),
   ('4.1','Unit & Integration Testing','Dev 2 (Junior)','36','Month 3'),
   ('4.2','User Acceptance Testing (UAT)','Both Devs','24','Month 3'),
   ('4.3','Bug Fixing & QA Review','Both Devs','32','Month 3'),
   ('5.1','Server Setup & Deployment','Dev 1 (Senior)','24','Month 3'),
   ('5.2','Documentation & User Manual','Dev 2 (Junior)','24','Month 3'),
   ('5.3','Final Handover & Training','Both Devs','12','Month 3'),
   ('TOTAL','---','---','708 hrs','3 Months')],
  widths=[0.6,2.6,1.3,0.7,0.8])

# 5. Cost Analysis
H(doc,'5. Cost Analysis')
B(doc,'Team: 2 Developers (1 Senior @ BDT 60,000/mo + 1 Junior @ BDT 30,000/mo) over 3 months. Bangladesh market rates, April 2026.')
H(doc,'5.1 Team & Labor Costs',size=12,color='2563EB',sb=6)
T(doc,['Role','Monthly Salary (BDT)','3-Month Cost (BDT)','Hours/Month'],
  [('Developer 1 - Senior / Full-Stack Lead','60,000','1,80,000','176 hrs'),
   ('Developer 2 - Junior / Frontend & QA','30,000','90,000','176 hrs'),
   ('TOTAL','90,000 / month','2,70,000','352 hrs / month')],
  widths=[2.5,1.6,1.6,1.5])
H(doc,'5.2 Direct Costs',size=12,color='2563EB',sb=6)
T(doc,['Category','Item','Total (BDT)'],
  [('Labor - Senior Dev','Full-stack development (3 months)','1,80,000'),
   ('Labor - Junior Dev','Frontend & QA (3 months)','90,000'),
   ('Design Tools','Figma Pro (3 months)','7,500'),
   ('AI Assist','GitHub Copilot x2 devs (3 months)','7,200'),
   ('Hardware','Laptop depreciation x2 devs','12,000'),
   ('Infrastructure','XAMPP, PHP, MySQL, Git (free/OSS)','0'),
   ('Testing','Device/browser testing allocation','3,000'),
   ('','DIRECT COST SUBTOTAL','2,99,700')],
  widths=[1.8,3.0,2.0])
H(doc,'5.3 Indirect Costs',size=12,color='2563EB',sb=6)
T(doc,['Category','Item','3-Month Total (BDT)'],
  [('Infrastructure','Internet - broadband 2 devs','4,200'),
   ('Infrastructure','Electricity / server allocation','3,000'),
   ('Administrative','Project management overhead','2,400'),
   ('Communication','Slack, cloud storage, email','1,500'),
   ('Training','Online references','1,500'),
   ('','INDIRECT COST SUBTOTAL','12,600')],
  widths=[1.8,3.0,2.0])
H(doc,'5.4 Cost Baseline & Total Budget',size=12,color='2563EB',sb=6)
T(doc,['Budget Component','Amount (BDT)','Notes'],
  [('A. Direct Costs','2,99,700','Labor, tools, hardware'),
   ('B. Indirect Costs','12,600','Overhead & operational'),
   ('COST BASELINE (A+B)','3,12,300','Approved performance reference'),
   ('C. Contingency Reserve (10%)','31,230','Known risks buffer'),
   ('CONTROL ACCOUNT BUDGET','3,43,530','With contingency'),
   ('D. Management Reserve (5%)','15,627','Unknown risks - sponsor held'),
   ('TOTAL PROJECT BUDGET','3,59,157','Full authorization level')],
  widths=[2.6,1.5,2.7])
H(doc,'5.5 Time-Phased Budget',size=12,color='2563EB',sb=6)
T(doc,['Cost Item','Month 1 (Apr)','Month 2 (May)','Month 3 (Jun)','Total'],
  [('Senior Developer','60,000','60,000','60,000','1,80,000'),
   ('Junior Developer','30,000','30,000','30,000','90,000'),
   ('Software Tools','4,900','4,900','4,900','14,700'),
   ('Hardware Depreciation','4,000','4,000','4,000','12,000'),
   ('Testing & Devices','1,000','1,000','1,000','3,000'),
   ('Indirect Costs','4,200','4,200','4,200','12,600'),
   ('Monthly Total','1,04,100','1,04,100','1,04,100','3,12,300'),
   ('Cumulative (Baseline)','1,04,100','2,08,200','3,12,300','3,12,300')],
  widths=[1.9,1.4,1.4,1.4,1.2])
H(doc,'5.6 Real-World vs Academic Cost Comparison',size=12,color='2563EB',sb=6)
T(doc,['Parameter','This Project (Academic)','Real-World Production'],
  [('Team Size','2 Developers','4-6 Developers (professional)'),
   ('Duration','3 Months','6-9 Months'),
   ('Senior Dev Rate','60,000/month','1,20,000-1,80,000/month'),
   ('Junior Dev Rate','30,000/month','60,000-80,000/month'),
   ('Hosting','XAMPP local (free)','AWS/DigitalOcean 8,000-15,000/mo'),
   ('Medicine Database','Manually curated (free)','IMS Health API 22,000-55,000/mo'),
   ('AI Engine','Knowledge-based (free)','OpenAI GPT-4o 3,300-8,800/mo'),
   ('Total Budget','3,59,157 BDT','20,00,000-50,00,000+ BDT')],
  widths=[2.0,2.2,2.6])
H(doc,'5.7 Earned Value Analysis (EVM) Framework',size=12,color='2563EB',sb=6)
T(doc,['EVM Metric','Formula','End-of-Month-1 Reference'],
  [('Budget at Completion (BAC)','Total cost baseline','3,12,300 BDT'),
   ('Planned Value (PV)','BAC x % work planned','1,04,100 BDT (33.3%)'),
   ('Cost Variance (CV)','EV - AC','+ve = under budget'),
   ('Schedule Variance (SV)','EV - PV','+ve = ahead of schedule'),
   ('Cost Performance Index (CPI)','EV / AC','>1 = cost-efficient'),
   ('Schedule Performance Index (SPI)','EV / PV','>1 = on/ahead of schedule'),
   ('Estimate at Completion (EAC)','BAC / CPI','Revised forecast total'),
   ('Variance at Completion (VAC)','BAC - EAC','+ve = under budget; -ve = over budget')],
  widths=[2.1,1.8,2.5])

# 6. SWOT
H(doc,'6. SWOT Analysis')
B(doc,'A SWOT Analysis evaluates the internal strengths and weaknesses of the Medicure platform alongside external opportunities and threats. This analysis is conducted from the perspective of a real-world health-tech product launch in Bangladesh.')

H(doc,'Strengths',size=12,color='059669',sb=6)
T(doc,['#','Strength & Detail','Competitive Advantage'],
  [('1','Largest Free BD-Specific Medicine Database\n\nMedicure contains 21,000+ Bangladeshi brand medicines curated specifically for the local drug market. No competitor platform in Bangladesh offers this level of localised, offline-accessible, free drug data. Users can search by brand name, generic name, side effects, and indications without needing an internet connection for cached results.','Unique market differentiator that cannot easily be replicated by global platforms unfamiliar with DGDA-approved local medicines.'),
   ('2','Truly All-in-One Health Management Platform\n\nUnlike single-purpose apps (alarm apps, drug databases, or pharmacy apps), Medicure integrates medication tracking, AI triage, drug search, family management, pharmacy ordering, adherence reports, health calculators, and a full admin panel into one seamless experience. No context-switching between multiple apps.','Reduces user friction significantly - everything a patient needs is in one authenticated session.'),
   ('3','Completely Free for All Patients\n\nMedicure operates on a zero-cost model for end users. There are no subscription fees, no premium tiers, and no paywalls. In a market like Bangladesh where disposable income for health apps is limited, this is a critical adoption driver. The admin-managed model keeps operational control without burdening patients financially.','Removes the #1 barrier to health app adoption in low-to-middle income markets.'),
   ('4','Works on Any Device - No App Store Required\n\nBuilt as a mobile-first responsive web app with Progressive Web App (PWA) support, Medicure can be installed directly from the browser on any Android or iOS device without going through the Google Play Store or Apple App Store. This bypasses app store approval delays, download friction, and storage concerns.','Instant reach to all smartphone users regardless of OS version or storage availability.'),
   ('5','Intelligent Browser-Based Alarm System\n\nMedicure implements a JavaScript-powered alarm system using Service Workers that fires browser notifications at the exact scheduled medication time - even when the tab is minimized or the phone screen is off. Users can mark doses as taken, snoozed, or missed directly from the notification, updating their adherence record in real time.','Solves the core medication non-adherence problem with a zero-dependency, zero-cost native browser solution.'),
   ('6','Comprehensive Admin Dashboard & Control Panel\n\nThe administrator has full visibility over every patient: their registered medicines, scheduled doses, adherence history, family members, and profile details. Admins can add/remove patients, send email broadcasts to all users, and monitor platform-wide activity. This makes Medicure deployable in clinic, hospital, or NGO settings where oversight is required.','Enables institutional deployment in healthcare settings beyond individual use.'),
   ('7','Printable 7-Day Adherence Report for Doctor Visits\n\nThe reports module generates a formatted, printable 7-day medication adherence summary showing which doses were taken, missed, or snoozed. Patients can bring this printout to doctor consultations, enabling physicians to assess treatment compliance and adjust prescriptions accordingly - closing a critical care gap.','Bridges the patient-doctor communication gap with tangible, standardised health records.'),
   ('8','Open-Source Stack - Zero Infrastructure Licensing Cost\n\nThe entire platform runs on PHP, MySQL, XAMPP, and PHPMailer - 100% free, open-source technologies. This eliminates licensing costs entirely and means any developer can contribute, extend, or deploy the platform without vendor lock-in. The codebase is public on GitHub, enabling community contributions and institutional adoption.','BDT 26,00,000+ annual savings vs enterprise alternatives; enables community-driven growth.')],
  widths=[0.3,3.8,2.7], hbg='059669')

H(doc,'Weaknesses',size=12,color='D97706',sb=6)
T(doc,['#','Weakness & Detail','Severity','Improvement Path'],
  [('1','Local Hosting Only - No Live Deployment\n\nMedicure currently runs on XAMPP (localhost), meaning it is only accessible on the developer\'s machine. There is no public URL, no live server, and no way for real patients outside the local network to access the platform. This is the single biggest barrier to real-world adoption.','Critical','Deploy to DigitalOcean Droplet or AWS Lightsail within one sprint. Cost: ~$12/month. Estimated setup time: 1-2 days.'),
   ('2','No Native Android / iOS Application\n\nWhile Medicure is PWA-installable, it lacks a true native app. Native apps support background push notifications, offline-first data sync, biometric authentication, and tighter OS integration. Without a Play Store listing, the platform is invisible to users browsing the app store for health management tools.','High','Build a React Native or Flutter wrapper around the existing web app. Reuse all existing API endpoints. Publish to Google Play Store BD. Estimated effort: 3-4 weeks.'),
   ('3','Email Broadcast Requires Manual SMTP Configuration\n\nThe admin email broadcast feature depends on PHPMailer with a Gmail SMTP account manually configured in a config file. This is not plug-and-play - it requires technical setup, is subject to Gmail sending limits (500 emails/day), and will fail silently if credentials expire or 2FA blocks the SMTP connection.','Medium','Migrate to a transactional email service like SendGrid (free tier: 100 emails/day) or AWS SES. Implement email delivery status tracking and bounce handling in the admin panel.'),
   ('4','No In-App Payment Gateway\n\nThe pharmacy Buy Medicine feature redirects users to Arogga\'s website with the medicine pre-searched via URL. No actual transaction happens within Medicure. This creates a disjointed user experience and means Medicure cannot earn transaction revenue or track purchase completion. Users may abandon the flow after being redirected.','Medium','Integrate bKash API or SSLCommerz (Bangladesh\'s leading payment gateway) for in-app checkout. Negotiate a commission structure with Arogga or Chaldal Health.'),
   ('5','Plain-Text Password Storage in Demo Mode\n\nFor development and admin setup purposes, some passwords may be stored or handled in plain text at certain points in the codebase. In production, all passwords must be hashed using bcrypt or Argon2. Storing plain-text passwords is a critical security vulnerability that could expose all user accounts if the database is breached.','Critical','Audit all password handling code. Enforce password_hash() / password_verify() across all authentication flows. Add a forced password change on first admin login.'),
   ('6','Single Admin Account - No Role-Based Access Control\n\nThe platform supports only one administrator account. There is no concept of Super Admin, Sub-Admin, or read-only Analyst roles. This limits scalability for clinic deployments where multiple staff members need different levels of access, and creates a single point of failure if admin credentials are compromised.','High','Design and implement a Role-Based Access Control (RBAC) system with at least three roles: Super Admin, Staff Admin, and Read-Only Analyst. Store roles in the database.'),
   ('7','No Bangla (Bengali) Language Support\n\nThe entire UI is in English, which limits accessibility for rural patients, elderly users, and individuals with lower English literacy - a significant portion of Bangladesh\'s 170M population. Health information is most effective when consumed in one\'s native language.','Medium','Implement an i18n (internationalization) framework. Create a Bengali language pack for all UI strings. Add a language toggle in the user profile settings page.'),
   ('8','No Caching Layer - Performance Bottleneck at Scale\n\nEvery page load triggers fresh database queries with no caching. For the 21,000+ medicines search, this means full table scans on every keystroke in the autocomplete. Under concurrent load (e.g. 100+ simultaneous users), this will cause significant latency and potential server timeout errors.','Medium','Implement Redis or Memcached for medicine search result caching. Add MySQL query result caching for frequently accessed static data. Implement pagination on all list views.')],
  widths=[0.3,3.5,0.7,2.3], hbg='D97706')
H(doc,'Opportunities',size=12,color='2563EB',sb=6)
T(doc,['#','Opportunity & Detail','Impact','Action Required'],
  [('1','Growing Digital Health Market in BD\n\nBangladesh has 170M+ population with rising smartphone penetration. The government Digital Bangladesh initiative and post-COVID demand have accelerated telehealth adoption. Medicure is perfectly positioned to capitalise on this.','High','Deploy to cloud; register under a2i / ICT Division BD'),
   ('2','Formal Pharmacy Partnership (Arogga, Chaldal Health)\n\nThe Buy Medicine feature currently redirects to Arogga via URL. A formal API or affiliate agreement could generate referral revenue per purchase, making the platform financially self-sustaining.','High','Negotiate affiliate/API deal with Arogga or Chaldal Health'),
   ('3','Doctor & Clinic Portal Expansion\n\nAdding a doctor-facing dashboard where physicians view patient medication history, adherence reports, and symptom logs would transform Medicure into a full care-coordination platform and multiply its market value.','High','Build doctor role with read-only patient report access and appointment booking'),
   ('4','Government eHealth Integration\n\nBangladesh Shasthya Batayon (16789) national helpline and DGDA publish official drug lists. Integration would boost credibility and provide verified, regularly updated medicine data directly from the regulator.','Medium','Apply for DGDA API access; align with a2i national eHealth framework'),
   ('5','Bangla Language Support\n\nThe UI is currently English-only. Full Bangla support would reach millions of rural and semi-urban users comfortable only in their native language, dramatically expanding the addressable user base beyond the educated urban segment.','High','Implement i18n with a Bengali language pack across all pages'),
   ('6','Cloud Deployment & Scalability\n\nMigrating from XAMPP localhost to DigitalOcean, AWS Lightsail, or Azure BD region makes the app accessible nationwide 24/7. A Cloudflare CDN layer could serve thousands of concurrent users with low latency.','High','Migrate to DigitalOcean Droplet 2GB RAM (~$12/mo); set up CI/CD'),
   ('7','AI Upgrade: GPT-Powered Symptom Checker\n\nThe current knowledge-based rule engine is limited. Upgrading to OpenAI GPT-4o or Google Gemini would handle complex symptom combinations, multi-disease differentials, and generate personalised health insights - making the product significantly more competitive.','Medium','Integrate OpenAI API with rate-limiter; keep mandatory disclaimer intact'),
   ('8','Native Mobile App (React Native / Flutter)\n\nDespite PWA support, a dedicated Google Play Store listing would dramatically increase discoverability, enable native push notifications for medicine reminders, and build stronger user trust through an app icon on the home screen.','High','Build React Native or Flutter wrapper; publish on Google Play Store BD')],
  widths=[0.3,3.5,0.7,2.3])

H(doc,'Threats & Risks',size=12,color='DC2626',sb=6)
T(doc,['#','Threat / Risk & Detail','Risk Level','Mitigation Strategy'],
  [('1','Competition from Established Platforms\n\nApps like Praava Health BD, Ada Health, Babylon, and WebMD have larger engineering teams, marketing budgets, and brand recognition. They can replicate Medicure features quickly and have existing user bases.',
    'Likelihood: High\nImpact: Medium',
    'Focus on the Bangladeshi drug market, Bangla language, and local pharmacy integrations - areas global platforms cannot easily replicate.'),
   ('2','Data Privacy & Healthcare Regulation Risk\n\nUpcoming BD data protection legislation (GDPR-modeled) will impose strict rules on health data storage. Sensitive data such as medications, symptoms, and family records stored without encryption creates serious legal liability.',
    'Likelihood: Medium\nImpact: High',
    'Implement HTTPS/SSL, add a Privacy Policy page, obtain explicit user consent on registration, and architect for GDPR-like compliance from day one.'),
   ('3','Patient Over-Reliance on AI Symptom Checker\n\nUsers may trust the AI Triage feature too heavily and delay seeking proper medical attention. This creates medical liability if a patient makes a harmful health decision based on the app output.',
    'Likelihood: Medium\nImpact: High',
    'Mandatory disclaimer on every AI response. Add a prominent See a Doctor CTA button. Implement severity flags that redirect critical symptoms to emergency contacts.'),
   ('4','Internet Dependency in Rural Bangladesh\n\nArogga redirect, AI checker, email notifications, and medicine search all require stable internet. In rural, char, and haor areas of Bangladesh, connectivity is unreliable and costly, excluding a large potential user base.',
    'Likelihood: High\nImpact: Medium',
    'Implement offline-first PWA caching using Service Workers and IndexedDB for medicine list, reminders, and recently viewed drug information.'),
   ('5','Medicine Database Staleness & Accuracy Risk\n\nThe 21,000+ BD medicine database was manually curated at a single point in time. DGDA approves new medicines monthly and withdraws others. Outdated data could seriously mislead users searching for discontinued drugs.',
    'Likelihood: High\nImpact: Medium',
    'Quarterly database review cycle. Apply for DGDA official drug list feed. Build an admin panel feature to flag, update, and retire individual medicine entries.'),
   ('6','Security Vulnerabilities - SQL Injection, XSS, CSRF\n\nThe platform stores sensitive health data and handles user authentication. Without proper input sanitization across all PHP endpoints, it is vulnerable to SQL injection, XSS, CSRF, and session hijacking attacks.',
    'Likelihood: Medium\nImpact: High',
    'Audit all queries to use PDO prepared statements. Add CSP headers, CSRF tokens, and rate limiting on login. Conduct a full security penetration test before any production launch.'),
   ('7','Single Admin Account - Critical Point of Failure\n\nOnly one admin account is supported. Lost or compromised credentials have no recovery path without direct database access. This is an unacceptable operational risk for any production healthcare deployment.',
    'Likelihood: Low\nImpact: High',
    'Build a multi-admin role system (Super Admin / Sub-Admin). Add 2FA via TOTP or SMS OTP for all admin logins. Implement an admin activity audit log for accountability.'),
   ('8','Scalability Failure Under High Traffic Load\n\nXAMPP/PHP with no caching layer cannot sustain high concurrent load. A media mention, government adoption, or viral event could cause a traffic spike that crashes the server and causes complete service unavailability.',
    'Likelihood: Low\nImpact: High',
    'Add Redis query caching and an Nginx reverse proxy. Deploy on a VPS with horizontal auto-scaling. Implement a CDN for static assets to reduce origin server load under peak demand.')],
  widths=[0.25,3.05,1.2,2.3], hbg='DC2626')

# 7. Conclusion
H(doc,'7. Conclusion')
B(doc,'Medicure is a well-rounded, practical, and locally-relevant digital health solution. Built by a 2-developer team over 3 months at a total project budget of BDT 3,59,157, it demonstrates strong command of full-stack web development, database design, UX principles, and real-world health technology challenges.')
B(doc,'By leveraging 100% free and open-source tools (PHP, MySQL, XAMPP), the project saved an estimated BDT 26,00,000+ per year compared to a fully commercial, enterprise-grade alternative. With further development - particularly live deployment, native mobile app development, and formal pharmacy partnerships - Medicure has genuine potential to become a meaningful healthcare tool for patients across Bangladesh.')

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Medicure - Final Project Report  |  CSE495 IT Project Management  |  East West University  |  April 2026')
r.italic=True; r.font.size=Pt(9); r.font.color.rgb=rgb('888888')

doc.save(OUT)
print(f'[SUCCESS] Saved: {OUT}')
