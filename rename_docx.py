import docx

def replace_text_in_runs(runs, words):
    for run in runs:
        for old, new in words.items():
            if old in run.text:
                run.text = run.text.replace(old, new)

def process_document(file_path, output_path):
    doc = docx.Document(file_path)
    words = {
        'MediCure': 'Medicature',
        'medicure': 'medicature',
        'Medicure': 'Medicature',
        'MEDICURE': 'MEDICATURE'
    }
    
    for para in doc.paragraphs:
        replace_text_in_runs(para.runs, words)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_runs(para.runs, words)
                    
    doc.save(output_path)
    print(f"Saved {output_path}")

process_document("1.docx", "1_Medicature.docx")
