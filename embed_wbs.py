"""
Inserts the WBS diagram image into Medicure_Final_Report.docx
in the WBS section (Section 4), replacing the existing text-only WBS table.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WBS_IMAGE = r'c:\xampp\htdocs\medicure\wbs_diagram.png'
REPORT    = r'c:\xampp\htdocs\medicure\Medicure_Final_Report.docx'

doc = Document(REPORT)

def rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

# ── Find the WBS section heading paragraph ────────────────────────────────────
wbs_para_idx = None
for i, para in enumerate(doc.paragraphs):
    if '4. Work Breakdown Structure' in para.text:
        wbs_para_idx = i
        break

if wbs_para_idx is None:
    print('[ERROR] Could not find WBS section in document.')
    exit(1)

print(f'[INFO] Found WBS section at paragraph index {wbs_para_idx}')

# ── Find paragraph after the WBS heading (skip the body description para) ────
# We want to insert the image after the first body paragraph in the WBS section
# and BEFORE the WBS table.
# Strategy: find the first table after the WBS paragraph and insert image before it.

# Get all block-level elements (paragraphs + tables) in order
from docx.oxml.ns import qn
body = doc.element.body

all_blocks = list(body)  # all XML children

# Find the WBS heading XML element
wbs_heading_elem = doc.paragraphs[wbs_para_idx]._element

# Walk forward from WBS heading, find the first table element
insert_before_elem = None
found_heading = False
for elem in all_blocks:
    if elem is wbs_heading_elem:
        found_heading = True
        continue
    if found_heading and elem.tag.endswith('}tbl'):
        insert_before_elem = elem
        break

# ── Build image paragraph ─────────────────────────────────────────────────────
from docx.oxml import OxmlElement
import copy

# Create a new paragraph with the image
img_para = OxmlElement('w:p')
img_pPr  = OxmlElement('w:pPr')
img_jc   = OxmlElement('w:jc')
img_jc.set(qn('w:val'), 'center')
img_pPr.append(img_jc)
img_para.append(img_pPr)

# Create a caption paragraph
cap_para = OxmlElement('w:p')
cap_pPr  = OxmlElement('w:pPr')
cap_jc2  = OxmlElement('w:jc')
cap_jc2.set(qn('w:val'), 'center')
cap_pPr.append(cap_jc2)
cap_para.append(cap_pPr)
cap_r = OxmlElement('w:r')
cap_rPr = OxmlElement('w:rPr')
cap_i = OxmlElement('w:i')
cap_sz = OxmlElement('w:sz')
cap_sz.set(qn('w:val'), '18')  # 9pt
cap_rPr.append(cap_i); cap_rPr.append(cap_sz)
cap_r.append(cap_rPr)
cap_t = OxmlElement('w:t')
cap_t.text = 'Figure 1: Medicure WBS – Phase-Based | 5 Phases | 18 Work Packages | 708 Total Hours'
cap_r.append(cap_t)
cap_para.append(cap_r)

# Blank para before image
blank_before = OxmlElement('w:p')
blank_after  = OxmlElement('w:p')

# Use python-docx's add_picture via a temporary document then extract the run XML
import tempfile, os
tmp_doc = Document()
tmp_sec = tmp_doc.sections[0]
tmp_para = tmp_doc.add_paragraph()
tmp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tmp_para.add_run()
run.add_picture(WBS_IMAGE, width=Inches(6.8))
img_run_xml = tmp_para._element  # this is a w:p containing the picture run

# Insert into the actual document body before the WBS table
if insert_before_elem is not None:
    body.insert(list(body).index(insert_before_elem), blank_before)
    body.insert(list(body).index(insert_before_elem), copy.deepcopy(img_run_xml))
    body.insert(list(body).index(insert_before_elem), cap_para)
    body.insert(list(body).index(insert_before_elem), blank_after)
    print('[INFO] WBS image inserted before the WBS table.')
else:
    # Fallback: append at the end of WBS heading paragraph
    wbs_heading_elem.addnext(copy.deepcopy(img_run_xml))
    print('[INFO] WBS image inserted after WBS heading (fallback).')

doc.save(REPORT)
print(f'[SUCCESS] Saved: {REPORT}')
