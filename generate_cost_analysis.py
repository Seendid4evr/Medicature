
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Helpers ──────────────────────────────────────────────────────────────────

def hex_color(hex_str):
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def set_cell_bg(cell, hex_str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_str.lstrip('#'))
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), kwargs.get('val', 'single'))
        border.set(qn('w:sz'), kwargs.get('sz', '6'))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get('color', '1A3C5E'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color='1A3C5E', space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = hex_color(color)
    run.font.size = Pt(18 - (level-1)*3)
    return p

def add_body(doc, text, italic=False, space_after=4):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    p.style.font.size = Pt(10.5)
    if italic:
        for run in p.runs:
            run.italic = True
    return p

def make_table(doc, headers, rows, col_widths=None, header_bg='1A3C5E', alt_bg='E8F0FE'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        set_cell_border(cell, color='FFFFFF')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = hex_color('FFFFFF')
        run.font.size = Pt(10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = alt_bg if r_idx % 2 == 1 else 'FFFFFF'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_border(cell, color='CCCCCC', sz='4')
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.RIGHT if c_idx > 0 and ('৳' in str(val) or '%' in str(val) or str(val).replace(',','').replace('.','').lstrip('-').isdigit()) else WD_ALIGN_PARAGRAPH.LEFT
            p.alignment = align
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

# ─── Main Document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ══════════════════════════════════════════════════════════════════════
# COVER / TITLE
# ══════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(24)
run = title_p.add_run('MEDICURE DIGITAL HEALTH PLATFORM')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = hex_color('1A3C5E')

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run('Project Cost Analysis Report')
run2.bold = True
run2.font.size = Pt(15)
run2.font.color.rgb = hex_color('2563EB')

detail_p = doc.add_paragraph()
detail_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = detail_p.add_run('Team: 2 Developers  |  Duration: 3 Months  |  Context: Bangladesh (BDT)')
run3.italic = True
run3.font.size = Pt(10.5)
run3.font.color.rgb = hex_color('555555')

doc.add_paragraph()
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
# PROJECT OVERVIEW TABLE (5×2)
# ══════════════════════════════════════════════════════════════════════
add_heading(doc, '1. Project Overview', level=1)

ov_table = doc.add_table(rows=5, cols=2)
ov_table.style = 'Table Grid'
ov_table.alignment = WD_TABLE_ALIGNMENT.CENTER
ov_data = [
    ('Project Name',    'Medicure – Digital Health Platform'),
    ('Team Size',       '2 Developers (1 Senior, 1 Junior)'),
    ('Project Duration','3 Months (April 2026 – June 2026)'),
    ('Technology Stack','PHP, MySQL, HTML/CSS/JavaScript, XAMPP'),
    ('Project Context', 'Academic + Real-world digital health use-case, Bangladesh'),
]
for i, (label, value) in enumerate(ov_data):
    row = ov_table.rows[i]
    bg = 'E8F0FE' if i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], '1A3C5E')
    set_cell_bg(row.cells[1], bg)
    set_cell_border(row.cells[0], color='FFFFFF', sz='4')
    set_cell_border(row.cells[1], color='CCCCCC', sz='4')
    lrun = row.cells[0].paragraphs[0].add_run(label)
    lrun.bold  = True
    lrun.font.color.rgb = hex_color('FFFFFF')
    lrun.font.size = Pt(10)
    vrun = row.cells[1].paragraphs[0].add_run(value)
    vrun.font.size = Pt(10)
for row in ov_table.rows:
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(4.5)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
# 2. TEAM & RATE STRUCTURE
# ══════════════════════════════════════════════════════════════════════
add_heading(doc, '2. Team & Rate Structure', level=1)
add_body(doc, 'The project is executed by a two-person development team over a period of three months. '
              'Monthly salary rates reflect mid-market rates for software developers in Dhaka, Bangladesh.')
make_table(doc,
    headers=['Role', 'Skill Level', 'Monthly Salary (BDT)', '3-Month Cost (BDT)', 'Working Hours/Month'],
    rows=[
        ('Developer 1', 'Senior / Full-Stack Lead', '৳ 60,000', '৳ 1,80,000', '176 hrs'),
        ('Developer 2', 'Junior / Frontend & QA',   '৳ 30,000', '৳  90,000', '176 hrs'),
        ('TOTAL', '—', '৳ 90,000 / month', '৳ 2,70,000', '352 hrs / month'),
    ],
    col_widths=[1.5, 2.0, 1.5, 1.7, 1.7],
)

# ══════════════════════════════════════════════════════════════════════
# 3. WORK BREAKDOWN STRUCTURE COST ALLOCATION
# ══════════════════════════════════════════════════════════════════════
add_heading(doc, '3. WBS-Based Cost Allocation (Direct Labor)', level=1)
add_body(doc, 'Project work is decomposed into five phases following a Phase-Based WBS. '
              'Hours are distributed across developers and phases to reflect realistic effort.')

wbs_rows = [
    # Phase, WBS Item, Dev1 hrs, Dev2 hrs, Total hrs, Rate BDT/hr, Cost BDT
    ('Phase 1: Initiation',  '1.1 Requirements & Scope Analysis',       12,  8, 20, 341, '৳ 6,818'),
    ('Phase 1: Initiation',  '1.2 Feasibility Study',                    8,  4, 12, 341, '৳ 4,091'),
    ('Phase 1: Initiation',  '1.3 Project Charter & Planning',           8,  4, 12, 341, '৳ 4,091'),
    ('Phase 2: Design',      '2.1 UI/UX Wireframes & Mockups',          12, 20, 32, 341, '৳10,909'),
    ('Phase 2: Design',      '2.2 Database Schema Design',              16,  8, 24, 341, '৳ 8,182'),
    ('Phase 2: Design',      '2.3 System Architecture Planning',        12,  4, 16, 341, '৳ 5,455'),
    ('Phase 3: Development', '3.1 Backend (PHP, Auth, DB, APIs)',        96, 40,136, 341, '৳46,364'),
    ('Phase 3: Development', '3.2 Frontend (HTML/CSS/JS, Dashboard)',    32, 80,112, 341, '৳38,182'),
    ('Phase 3: Development', '3.3 Symptom Checker AI Module',           40, 16, 56, 341, '৳19,091'),
    ('Phase 3: Development', '3.4 Pharmacy & Medicines Module',         32, 24, 56, 341, '৳19,091'),
    ('Phase 3: Development', '3.5 Family Health Management Module',     24, 16, 40, 341, '৳13,636'),
    ('Phase 3: Development', '3.6 Reports & Health Calculators',        20, 20, 40, 341, '৳13,636'),
    ('Phase 4: Testing',     '4.1 Unit & Integration Testing',          12, 24, 36, 341, '৳12,273'),
    ('Phase 4: Testing',     '4.2 User Acceptance Testing (UAT)',        8, 16, 24, 341, '৳ 8,182'),
    ('Phase 4: Testing',     '4.3 Bug Fixing & QA',                     16, 16, 32, 341, '৳10,909'),
    ('Phase 5: Deployment',  '5.1 Server Setup & Deployment',           16,  8, 24, 341, '৳ 8,182'),
    ('Phase 5: Deployment',  '5.2 Documentation & User Manual',          8, 16, 24, 341, '৳ 8,182'),
    ('Phase 5: Deployment',  '5.3 Final Handover & Training',            4,  8, 12, 341, '৳ 4,091'),
    ('TOTAL', '—', 376, 332, 708, '—', '৳2,41,364'),
]

make_table(doc,
    headers=['Phase', 'WBS Task', 'Dev 1 (hrs)', 'Dev 2 (hrs)', 'Total (hrs)', 'Rate (BDT/hr)', 'Labor Cost (BDT)'],
    rows=wbs_rows,
    col_widths=[1.35, 2.2, 0.75, 0.75, 0.75, 0.85, 1.0],
)
add_body(doc, '* Blended hourly rate = (৳60,000 + ৳30,000) ÷ (2 × 132 working hrs/month*) ≈ ৳341/hr  |  *Assumes 22 working days × 6 hrs billable/day', italic=True)

# ══════════════════════════════════════════════════════════════════════
# 4. DIRECT & INDIRECT COSTS
# ══════════════════════════════════════════════════════════════════════
add_heading(doc, '4. Direct & Indirect Costs Breakdown', level=1)

add_heading(doc, '4.1 Direct Costs', level=2, color='2563EB', space_before=6)
make_table(doc,
    headers=['Cost Category', 'Item', 'Qty', 'Unit Cost (BDT)', 'Total (BDT)'],
    rows=[
        ('Labor – Senior Developer',  'Full-stack development (3 months)',       1, '৳60,000/mo', '৳1,80,000'),
        ('Labor – Junior Developer',  'Frontend & QA (3 months)',                1, '৳30,000/mo', '৳ 90,000'),
        ('Software & Tools',          'VS Code, Git, XAMPP, MySQL Workbench',    1, 'Free/OSS',   '৳       0'),
        ('Software & Tools',          'Adobe XD / Figma (Pro plan, 3 months)',   1, '৳2,500/mo',  '৳  7,500'),
        ('Software & Tools',          'GitHub Copilot subscription (3 months)',  2, '৳1,200/mo',  '৳  7,200'),
        ('Hardware / Equipment',      'Laptop usage allocation (depreciation)',   2, '৳2,000/mo',  '৳ 12,000'),
        ('Testing',                   'Device testing (Android/iOS simulators)', 1, 'Lump sum',   '৳  3,000'),
        ('',                          'DIRECT COST SUBTOTAL',                   '', '',            '৳2,99,700'),
    ],
    col_widths=[1.7, 2.4, 0.5, 1.2, 1.0],
)

add_heading(doc, '4.2 Indirect Costs', level=2, color='2563EB', space_before=6)
make_table(doc,
    headers=['Cost Category', 'Item', 'Monthly (BDT)', 'Total 3 Months (BDT)'],
    rows=[
        ('Infrastructure', 'Shared hosting / local server electricity',      '৳  1,000', '৳  3,000'),
        ('Infrastructure', 'Internet (broadband 2 devs)',                    '৳  1,400', '৳  4,200'),
        ('Administrative', 'Project management / meeting overhead',          '৳    800', '৳  2,400'),
        ('Communication',  'Slack, email, cloud storage (Google Drive)',     '৳    500', '৳  1,500'),
        ('Training',       'Online courses / documentation references',      '৳    500', '৳  1,500'),
        ('',               'INDIRECT COST SUBTOTAL',                        '',          '৳ 12,600'),
    ],
    col_widths=[1.7, 2.8, 1.2, 1.5],
)

# ══════════════════════════════════════════════════════════════════════
# 5. OPPORTUNITY & RISK-RELATED COSTS
# ══════════════════════════════════════════════════════════════════════
add_heading(doc, '5. Opportunity & Risk-Related Costs', level=1)
make_table(doc,
    headers=['Cost Type', 'Description', 'Estimated Amount (BDT)'],
    rows=[
        ('Opportunity Cost',   'Revenue/freelance income forgone by 2 devs for 3 months (~40% of salary)', '৳ 1,08,000'),
        ('Risk Contingency',   '10% of total project cost (Direct + Indirect) reserved for scope creep',   '৳  31,230'),
        ('Risk – Technical',   'Unexpected bug resolution, library upgrades, server failures',              '৳  10,000'),
        ('Risk – Scope Creep', 'Additional feature requests beyond original WBS',                          '৳  15,000'),
        ('',                   'RISK & OPPORTUNITY SUBTOTAL',                                              '৳  56,230'),
    ],
    col_widths=[1.6, 3.8, 1.4],
)

# ══════════════════════════════════════════════════════════════════════
# 6. COMPLETE COST BASELINE (BUDGET)
# ══════════════════════════════════════════════════════════════════════
add_heading(doc, '6. Cost Baseline & Total Project Budget', level=1)
add_body(doc, 'The Cost Baseline is the approved planned budget (excluding management reserves). '
              'It serves as the reference for Earned Value Analysis throughout the project lifecycle.')

make_table(doc,
    headers=['Budget Component', 'Amount (BDT)', 'Notes'],
    rows=[
        ('A. Direct Costs (Labor + Tools + Hardware)', '৳ 2,99,700', 'Primary project expenditure'),
        ('B. Indirect Costs (Overhead)',               '৳  12,600',  'Operational support costs'),
        ('── COST BASELINE (A + B)',                   '৳ 3,12,300', 'Approved performance reference'),
        ('C. Contingency Reserve (10%)',               '৳  31,230',  'Known risks buffer'),
        ('── CONTROL ACCOUNT BUDGET (A+B+C)',          '৳ 3,43,530', 'Project budget with contingency'),
        ('D. Management Reserve (5%)',                 '৳  15,627',  'Unknown unknowns – separate'),
        ('══ TOTAL PROJECT BUDGET (A+B+C+D)',          '৳ 3,59,157', 'Full authorization level'),
    ],
    col_widths=[2.8, 1.5, 2.5],
    header_bg='1A3C5E',
)

# ══════════════════════════════════════════════════════════════════════
# 7. TIME-PHASED BUDGET (MONTHLY DISTRIBUTION)
# ══════════════════════════════════════════════════════════════════════
add_heading(doc, '7. Time-Phased Budget (Monthly Cost Baseline)', level=1)
add_body(doc, 'Costs are distributed across three months based on planned work intensity per phase. '
              'Month 1 covers initiation and design; Month 2 is the heaviest development sprint; Month 3 covers testing, deployment, and wrap-up.')

make_table(doc,
    headers=['Cost Item', 'Month 1 – April 2026 (BDT)', 'Month 2 – May 2026 (BDT)', 'Month 3 – June 2026 (BDT)', 'Total (BDT)'],
    rows=[
        ('Senior Developer Labor',  '৳ 60,000', '৳ 60,000', '৳ 60,000', '৳ 1,80,000'),
        ('Junior Developer Labor',  '৳ 30,000', '৳ 30,000', '৳ 30,000', '৳  90,000'),
        ('Software Tools (Figma etc.)', '৳ 4,900', '৳ 4,900', '৳ 4,900', '৳  14,700'),
        ('Hardware Depreciation',   '৳ 4,000', '৳ 4,000', '৳ 4,000', '৳  12,000'),
        ('Testing / Devices',       '৳ 1,000', '৳ 1,000', '৳ 1,000', '৳   3,000'),
        ('Indirect Costs',          '৳ 4,200', '৳ 4,200', '৳ 4,200', '৳  12,600'),
        ('──── Monthly Total',      '৳1,04,100', '৳1,04,100', '৳1,04,100', '৳3,12,300'),
        ('Cumulative (Cost Baseline)', '৳1,04,100', '৳2,08,200', '৳3,12,300', '৳3,12,300'),
    ],
    col_widths=[2.0, 1.5, 1.5, 1.5, 1.3],
)
add_body(doc, 'Note: The even monthly distribution reflects fixed salaries. Variable sprint costs remain absorbed within overhead.', italic=True)

# ══════════════════════════════════════════════════════════════════════
# 8. EARNED VALUE ANALYSIS (EVM) PLAN
# ══════════════════════════════════════════════════════════════════════
add_heading(doc, '8. Earned Value Analysis (EVM) Framework', level=1)
add_body(doc, 'The following EVM parameters are defined for mid-project performance tracking (end of Month 1 snapshot).')

make_table(doc,
    headers=['EVM Term', 'Formula', 'Planned Value (End of Month 1)'],
    rows=[
        ('Budget at Completion (BAC)',          'Total cost baseline',              '৳ 3,12,300'),
        ('Planned Value (PV)',                  'BAC × % work planned',             '৳ 1,04,100  (33.3%)'),
        ('Earned Value (EV)',                   'BAC × % work completed',           'Measured at checkpoint'),
        ('Actual Cost (AC)',                    'Actual money spent',               'Tracked monthly'),
        ('Cost Variance (CV)',                  'EV − AC',                          '+ve = under budget'),
        ('Schedule Variance (SV)',              'EV − PV',                          '+ve = ahead of schedule'),
        ('Cost Performance Index (CPI)',        'EV ÷ AC',                          '>1 = cost-efficient'),
        ('Schedule Performance Index (SPI)',    'EV ÷ PV',                          '>1 = ahead of schedule'),
        ('Estimate at Completion (EAC)',        'BAC ÷ CPI',                        'Revised forecast'),
        ('Estimate to Complete (ETC)',          'EAC − AC',                         'Remaining cost needed'),
        ('Variance at Completion (VAC)',        'BAC − EAC',                        '+ve = under; −ve = over'),
    ],
    col_widths=[2.0, 1.8, 2.7],
)

# ══════════════════════════════════════════════════════════════════════
# 9. COST COMPARISON (REAL-WORLD VS. ACADEMIC)
# ══════════════════════════════════════════════════════════════════════
add_heading(doc, '9. Real-World vs. Academic Production Cost Comparison', level=1)
add_body(doc, 'This section benchmarks the Medicure academic project cost against a professionally contracted equivalent.')

make_table(doc,
    headers=['Parameter', 'Academic Project (This Report)', 'Real-World Production Estimate'],
    rows=[
        ('Team Size',               '2 Developers (Student/Junior rate)', '4–6 Developers (Professional rate)'),
        ('Duration',                '3 Months',                           '6–9 Months'),
        ('Senior Dev Rate/Month',   '৳ 60,000',                           '৳ 1,20,000 – ৳ 1,80,000'),
        ('Junior Dev Rate/Month',   '৳ 30,000',                           '৳ 60,000 – ৳ 80,000'),
        ('Labor Cost',              '৳ 2,70,000',                         '৳ 15,00,000 – ৳ 40,00,000'),
        ('Software & Tools',        '৳ 14,700 (OSS-based)',               '৳ 1,50,000+ (Enterprise licenses)'),
        ('Infrastructure/Hosting',  '৳ 3,000 (local)',                    '৳ 50,000+ (Cloud: AWS/Azure)'),
        ('QA & Security Testing',   'Included in dev hours',              '৳ 3,00,000 – ৳ 5,00,000'),
        ('Total Budget',            '৳ 3,59,157',                         '৳ 20,00,000 – ৳ 50,00,000+'),
    ],
    col_widths=[2.0, 2.25, 2.25],
)

# ══════════════════════════════════════════════════════════════════════
# 10. KEY ASSUMPTIONS & CONSTRAINTS
# ══════════════════════════════════════════════════════════════════════
add_heading(doc, '10. Key Assumptions & Constraints', level=1)
assumptions = [
    '1. Salary rates are based on the Dhaka, Bangladesh mid-market (2025–2026 reference).',
    '2. Both developers work full-time (approximately 6 billable hours/working day, 22 days/month).',
    '3. Development uses only open-source infrastructure (XAMPP, PHP, MySQL) — no cloud hosting costs incurred during development.',
    '4. No physical office space is required; all work is remote/home-based (utilities absorbed in indirect costs).',
    '5. The cost baseline is fixed at ৳3,12,300. Changes require a formal change request.',
    '6. Contingency reserve (৳31,230) is controlled by the project lead and released only upon approved risk events.',
    '7. Management reserve (৳15,627) is held by the project sponsor and not accessible without escalation.',
    '8. All costs are denominated in BDT (Bangladeshi Taka). No foreign exchange exposure.',
]
for a in assumptions:
    p = doc.add_paragraph(a, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()

# Closing note
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run('Medicure Project — Cost Analysis Report  |  CSE495 IT Project Management  |  East West University')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = hex_color('888888')

# ─── Save ─────────────────────────────────────────────────────────────────────
output_path = r'c:\xampp\htdocs\medicure\Medicure_Cost_Analysis.docx'
doc.save(output_path)
print(f'[SUCCESS] Saved to: {output_path}')
